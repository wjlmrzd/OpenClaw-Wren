# -*- coding: utf-8 -*-
"""
Word Document Generator - Python backend using python-docx
"""
import sys
import os
import json
import argparse
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("ERROR: python-docx not installed. Run: pip install python-docx")
    sys.exit(1)


def create_document(title="New Document", output="output.docx", template=None):
    """Create a new Word document."""
    if template and os.path.exists(template):
        doc = Document(template)
    else:
        doc = Document()
    
    # Add title
    if title:
        heading = doc.add_heading(title, 0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.save(output)
    return output


def add_paragraph(doc_path, text, style="Normal", bold=False, italic=False, align="left"):
    """Add a paragraph to an existing document."""
    doc = Document(doc_path)
    para = doc.add_paragraph(text, style=style)
    
    if bold or italic:
        for run in para.runs:
            run.bold = bold
            run.italic = italic
    
    # Set alignment
    align_map = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY
    }
    para.alignment = align_map.get(align, WD_ALIGN_PARAGRAPH.LEFT)
    
    doc.save(doc_path)
    return doc_path


def add_heading(doc_path, text, level=1):
    """Add a heading to the document."""
    doc = Document(doc_path)
    doc.add_heading(text, level=level)
    doc.save(doc_path)
    return doc_path


def add_table(doc_path, data, headers=None, output=None):
    """Add a table to the document."""
    doc = Document(doc_path) if os.path.exists(doc_path) else Document()
    
    if headers:
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = 'Light Grid Accent 1'
        hdr_cells = table.rows[0].cells
        for i, h in enumerate(headers):
            hdr_cells[i].text = str(h)
            # Make header bold
            for paragraph in hdr_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.bold = True
    
    if data:
        for row_data in data:
            row = table.add_row().cells
            for i, cell_data in enumerate(row_data):
                row[i].text = str(cell_data)
    
    output_path = output or doc_path
    doc.save(output_path)
    return output_path


def find_replace(doc_path, find_text, replace_text, output=None):
    """Find and replace text in document."""
    doc = Document(doc_path)
    count = 0
    
    for para in doc.paragraphs:
        if find_text in para.text:
            inline = para.runs
            for i, run in enumerate(inline):
                if find_text in run.text:
                    run.text = run.text.replace(find_text, replace_text)
                    count += 1
    
    # Also search in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if find_text in para.text:
                        for run in para.runs:
                            if find_text in run.text:
                                run.text = run.text.replace(find_text, replace_text)
                                count += 1
    
    output_path = output or doc_path
    doc.save(output_path)
    return output_path, count


def extract_text(doc_path, output=None):
    """Extract all text from document."""
    doc = Document(doc_path)
    text_parts = []
    
    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text)
    
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_text:
                text_parts.append(" | ".join(row_text))
    
    full_text = "\n".join(text_parts)
    
    if output:
        with open(output, "w", encoding="utf-8") as f:
            f.write(full_text)
        return output
    
    return full_text


def set_style(doc_path, style_name, font_name="Arial", font_size=12, bold=False):
    """Apply style to document."""
    doc = Document(doc_path)
    style = doc.styles[style_name]
    style.font.name = font_name
    style.font.size = Pt(font_size)
    style.font.bold = bold
    doc.save(doc_path)
    return doc_path


def main():
    parser = argparse.ArgumentParser(description="Word Document Generator")
    parser.add_argument("action", choices=["create", "add-paragraph", "add-heading", 
                                            "add-table", "find-replace", "extract"])
    parser.add_argument("--input", "-i")
    parser.add_argument("--output", "-o")
    parser.add_argument("--title", "-t")
    parser.add_argument("--template")
    parser.add_argument("--text")
    parser.add_argument("--find")
    parser.add_argument("--replace")
    parser.add_argument("--style", default="Normal")
    parser.add_argument("--level", type=int, default=1)
    parser.add_argument("--headers", nargs="+")
    parser.add_argument("--data", nargs="+")
    parser.add_argument("--align", default="left")
    parser.add_argument("--bold", action="store_true")
    parser.add_argument("--italic", action="store_true")
    
    args = parser.parse_args()
    
    if args.action == "create":
        result = create_document(args.title or "New Document", 
                                  args.output or "output.docx",
                                  args.template)
        print(f"Created: {result}")
    
    elif args.action == "add-paragraph":
        result = add_paragraph(args.input, args.text or "", args.style,
                               args.bold, args.italic, args.align)
        print(f"Added paragraph to: {result}")
    
    elif args.action == "add-heading":
        result = add_heading(args.input, args.text or "", args.level)
        print(f"Added heading to: {result}")
    
    elif args.action == "add-table":
        data = [args.data] if args.data else []
        headers = args.headers
        result = add_table(args.input or "output.docx", data, headers, args.output)
        print(f"Added table to: {result}")
    
    elif args.action == "find-replace":
        result, count = find_replace(args.input, args.find, args.replace, args.output)
        print(f"Replaced {count} occurrences in: {result}")
    
    elif args.action == "extract":
        result = extract_text(args.input, args.output)
        if args.output:
            print(f"Extracted to: {result}")
        else:
            print(result)


if __name__ == "__main__":
    main()
